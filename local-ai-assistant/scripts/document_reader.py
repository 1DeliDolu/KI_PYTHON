"""
Document Reader Module
Handles reading and processing various document formats
"""

import os
import logging
from typing import Dict, Optional, Union
import PyPDF2
from docx import Document
import chardet
import sys

# Add parent directory to path for importing config
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import Config

class DocumentReader:
    def __init__(self):
        self.config = Config()
        self.logger = logging.getLogger(__name__)
    
    def read_document(self, file_path: str) -> Dict:
        """
        Read document content based on file extension
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Dict: Document content and metadata
        """
        try:
            # Validate file
            is_valid, message = self.config.validate_file(file_path, "document")
            if not is_valid:
                return {
                    "success": False,
                    "message": message,
                    "content": ""
                }
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                content = self._read_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                content = self._read_docx(file_path)
            elif file_ext == '.txt':
                content = self._read_text(file_path)
            elif file_ext == '.rtf':
                content = self._read_rtf(file_path)
            else:
                return {
                    "success": False,
                    "message": f"Unsupported file format: {file_ext}",
                    "content": ""
                }
            
            if content is None:
                return {
                    "success": False,
                    "message": "Failed to read document content",
                    "content": ""
                }
            
            # Get file metadata
            metadata = self._get_file_metadata(file_path)
            
            return {
                "success": True,
                "file_path": file_path,
                "file_type": file_ext,
                "content": content,
                "metadata": metadata,
                "word_count": len(content.split()),
                "character_count": len(content)
            }
            
        except Exception as e:
            self.logger.error(f"Error reading document '{file_path}': {e}")
            return {
                "success": False,
                "message": f"Error reading document: {str(e)}",
                "content": ""
            }
    
    def _read_pdf(self, file_path: str) -> Optional[str]:
        """Read PDF file content"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"
                
                return content.strip()
                
        except Exception as e:
            self.logger.error(f"Error reading PDF '{file_path}': {e}")
            return None
    
    def _read_docx(self, file_path: str) -> Optional[str]:
        """Read DOCX file content"""
        try:
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Also read tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    content += row_text + "\n"
            
            return content.strip()
            
        except Exception as e:
            self.logger.error(f"Error reading DOCX '{file_path}': {e}")
            return None
    
    def _read_text(self, file_path: str) -> Optional[str]:
        """Read plain text file content"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding = chardet.detect(raw_data)['encoding']
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding or 'utf-8') as file:
                return file.read()
                
        except Exception as e:
            self.logger.error(f"Error reading text file '{file_path}': {e}")
            return None
    
    def _read_rtf(self, file_path: str) -> Optional[str]:
        """Read RTF file content (basic implementation)"""
        try:
            # For RTF files, we'll read as text and do basic cleanup
            # This is a simplified approach - for full RTF support, use striprtf library
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
                # Basic RTF cleanup (remove common RTF tags)
                import re
                content = re.sub(r'\\[a-z]+\d*\s?', '', content)  # Remove RTF commands
                content = re.sub(r'[{}]', '', content)  # Remove braces
                content = re.sub(r'\s+', ' ', content)  # Normalize whitespace
                
                return content.strip()
                
        except Exception as e:
            self.logger.error(f"Error reading RTF '{file_path}': {e}")
            return None
    
    def _get_file_metadata(self, file_path: str) -> Dict:
        """Get file metadata"""
        try:
            stat = os.stat(file_path)
            
            return {
                "file_name": os.path.basename(file_path),
                "file_size_bytes": stat.st_size,
                "file_size_mb": round(stat.st_size / (1024 * 1024), 2),
                "created_time": stat.st_ctime,
                "modified_time": stat.st_mtime,
                "accessed_time": stat.st_atime
            }
            
        except Exception as e:
            self.logger.error(f"Error getting metadata for '{file_path}': {e}")
            return {}
    
    def extract_text_summary(self, content: str, max_sentences: int = 3) -> str:
        """
        Extract a summary from document content
        
        Args:
            content (str): Document content
            max_sentences (int): Maximum number of sentences in summary
            
        Returns:
            str: Text summary
        """
        try:
            sentences = content.split('. ')
            if len(sentences) <= max_sentences:
                return content
            
            # Simple extractive summary - take first few sentences
            summary_sentences = sentences[:max_sentences]
            summary = '. '.join(summary_sentences)
            
            if not summary.endswith('.'):
                summary += '.'
            
            return summary
            
        except Exception as e:
            self.logger.error(f"Error creating summary: {e}")
            return content[:500] + "..." if len(content) > 500 else content
    
    def search_in_document(self, content: str, search_term: str, case_sensitive: bool = False) -> Dict:
        """
        Search for a term in document content
        
        Args:
            content (str): Document content
            search_term (str): Term to search for
            case_sensitive (bool): Whether search should be case sensitive
            
        Returns:
            Dict: Search results with occurrences and context
        """
        try:
            if not case_sensitive:
                content_search = content.lower()
                term_search = search_term.lower()
            else:
                content_search = content
                term_search = search_term
            
            occurrences = []
            start = 0
            
            while True:
                index = content_search.find(term_search, start)
                if index == -1:
                    break
                
                # Get context around the found term
                context_start = max(0, index - 50)
                context_end = min(len(content), index + len(search_term) + 50)
                context = content[context_start:context_end]
                
                occurrences.append({
                    "index": index,
                    "context": context.strip(),
                    "line_number": content[:index].count('\n') + 1
                })
                
                start = index + 1
            
            return {
                "search_term": search_term,
                "total_occurrences": len(occurrences),
                "case_sensitive": case_sensitive,
                "occurrences": occurrences
            }
            
        except Exception as e:
            self.logger.error(f"Error searching in document: {e}")
            return {
                "search_term": search_term,
                "total_occurrences": 0,
                "case_sensitive": case_sensitive,
                "occurrences": [],
                "error": str(e)
            }